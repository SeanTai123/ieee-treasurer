import streamlit as st
import gspread
from google.oauth2.service_account import Credentials
import pandas as pd
from datetime import datetime
import io
import requests
import base64
from docx import Document
from docx.shared import Mm
import plotly.express as px
import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart

# --- 1. CONFIGURATION ---
st.set_page_config(page_title="IEEE Treasury App", page_icon="💸", layout="wide")

SHEET_URL = "https://docs.google.com/spreadsheets/d/1kBFKNBkSNLJS4-qJrZ1QfhRGMkyGJCEft0T89JreTXw/edit?usp=sharing" 
DRIVE_FOLDER_ID = "1_SbBPQE-kYHhiate_sLUiPWJ9ypFIaBZ"
IMGBB_API_KEY = "8d854f84a4a2b638d00b89f0deca599b"

# --- 2. EMAIL HELPER FUNCTION (NOTIFY ADMIN) ---
def notify_admin_email(payee: str, txn_id: str, amount: float, date: str, description: str, category: str):
    """Send a notification email to the Treasurer when a new claim is submitted."""
    try:
        sender_email = st.secrets["email"]["sender"]
        admin_email = st.secrets["email"]["admin_email"]
        app_password = st.secrets["email"]["password"]
        
        msg = MIMEMultipart("alternative")
        msg['Subject'] = f"🚨 New Claim Submitted: {txn_id} by {payee}"
        msg['From']    = sender_email
        msg['To']      = admin_email

        body = f"""
        <html><body>
        <h3>New Claim Requires Approval</h3>
        <p>A new expense claim has been submitted on the Treasury App.</p>
        <table border="1" cellpadding="6" cellspacing="0">
            <tr><td><b>Claimant</b></td><td>{payee}</td></tr>
            <tr><td><b>Amount</b></td><td><b>RM {amount:.2f}</b></td></tr>
            <tr><td><b>Category</b></td><td>{category}</td></tr>
            <tr><td><b>Description</b></td><td>{description}</td></tr>
            <tr><td><b>Date</b></td><td>{date}</td></tr>
            <tr><td><b>Transaction ID</b></td><td>{txn_id}</td></tr>
        </table>
        <p>Please log in to the App's Admin Portal (Tab 7: Approvals) to review it.</p>
        </body></html>
        """
        msg.attach(MIMEText(body, "html"))

        with smtplib.SMTP_SSL("smtp.gmail.com", 465) as server:
            server.login(sender_email, app_password)
            server.sendmail(sender_email, admin_email, msg.as_string())

        return True
    except Exception as e:
        st.warning(f"⚠️ App updated, but email notification failed: {e}")
        return False

# ==========================================
# 3. AUTH GATE (SIDEBAR)
# ==========================================
if 'admin_auth' not in st.session_state:
    st.session_state['admin_auth'] = False

with st.sidebar:
    st.header("🔒 Admin Portal")
    if not st.session_state['admin_auth']:
        st.markdown("Login to access Treasury Ledgers and Management Tools.")
        pwd = st.text_input("Enter Admin Password", type="password")
        if st.button("Login"):
            correct_pwd = st.secrets.get("admin_password", "ieee2026") 
            if pwd == correct_pwd:
                st.session_state['admin_auth'] = True
                st.rerun()
            else:
                st.error("Incorrect password.")
    else:
        st.success("✅ Logged in as Admin")
        if st.button("Logout"):
            st.session_state['admin_auth'] = False
            st.rerun()

IS_TREASURER = st.session_state['admin_auth']

# --- 4. CONNECT TO SHEETS ---
scopes = [
    "https://www.googleapis.com/auth/spreadsheets",
    "https://www.googleapis.com/auth/drive"
]
credentials = Credentials.from_service_account_info(
    st.secrets["gcp_service_account"],
    scopes=scopes
)
client = gspread.authorize(credentials)
sheet = client.open_by_url(SHEET_URL).sheet1

st.title("💸 IEEE Treasury Management")

# Setup Tabs based on role
if IS_TREASURER:
    tabs = st.tabs([
        "📝 Submit Claim", "📊 Monthly Ledger", "📄 Word Tracker",
        "📈 Dashboard", "⏳ Pending Claims", "🛠️ Manage", "🔍 Approvals"
    ])
    tab1, tab2, tab3, tab4, tab5, tab6, tab7 = tabs
else:
    tabs = st.tabs(["📝 Submit Claim"])
    tab1 = tabs[0]

# ==========================================
# TAB 1: CLAIM SUBMISSION (PUBLIC)
# ==========================================
with tab1:
    st.markdown("Submit expenses. Receipts are saved to ImgBB and linked automatically.")
    with st.form("claim_form", clear_on_submit=True):
        date = st.date_input("Date of Transaction")
        payee = st.text_input("Your Name (Claimant / Payer)")
        description = st.text_input("Description (e.g., Paper plates for Welcoming Night)")
        
        category = st.selectbox("Category", ["Event Expense", "Merch Sales", "Operations", "Sponsorship", "Other"])
        other_category = st.text_input("If Category is 'Other', please specify the event/details:")
        
        txn_type = st.radio("Transaction Type", ["Expense (Claim)", "Income (Revenue)"])
        amount = st.number_input("Amount (RM)", min_value=0.0, format="%.2f")
        
        receipt_file = st.file_uploader("Upload Receipt Image (PNG/JPG)", type=['png', 'jpg', 'jpeg'])
        submit = st.form_submit_button("Submit Transaction")

        if submit:
            if not payee or not description or amount <= 0:
                st.error("Please fill in all details and ensure the amount is greater than 0.")
            elif category == "Other" and not other_category.strip():
                st.error("🚨 You selected 'Other' for the category. Please specify what it is for in the text box!")
            else:
                with st.spinner("Processing file and updating ledger..."):
                    records = sheet.get_all_records()
                    if len(records) > 0:
                        last_record = records[-1]
                        last_balance_str = str(last_record.get('Running Balance', '0')).replace('RM', '').replace(',', '').strip()
                        last_balance = float(last_balance_str) if last_balance_str else 0.0
                        
                        last_id = last_record.get('Transaction ID', 'TRX-0000')
                        try:
                            id_num = int(last_id.split('-')[1])
                            new_id = f"TRX-{id_num + 1:04d}"
                        except:
                            new_id = f"TRX-{len(records) + 1:04d}"
                    else:
                        last_balance = 0.0
                        new_id = "TRX-0001"

                    file_id = "No Receipt"
                    if receipt_file is not None:
                        try:
                            b64_image = base64.b64encode(receipt_file.getvalue()).decode('utf-8')
                            res = requests.post(
                                "https://api.imgbb.com/1/upload",
                                data={"key": IMGBB_API_KEY, "image": b64_image, "name": f"{new_id}_{payee}"}
                            )
                            if res.status_code == 200:
                                file_id = res.json()['data']['url']
                            else:
                                st.error(f"Image upload failed: {res.text}")
                                st.stop()
                        except Exception as e:
                            st.error(f"🚨 Upload Failed! Error: {str(e)}")
                            st.stop()

                    is_income = "Income" in txn_type
                    income_val = amount if is_income else ""
                    expense_val = amount if not is_income else ""
                    new_balance = (last_balance + amount) if is_income else (last_balance - amount)

                    final_category = other_category.strip() if category == "Other" else category

                    row_data = [
                        new_id, str(date), description, final_category,
                        f"{income_val:.2f}" if income_val else "",
                        f"{expense_val:.2f}" if expense_val else "",
                        file_id, payee,
                        "Cleared" if is_income else "Awaiting Approval",
                        "", f"RM {new_balance:,.2f}"
                    ]
                    
                    sheet.append_row(row_data)
                    
                    # Send Email Notification to Admin if it's an expense claim
                    if not is_income:
                        notify_admin_email(
                            payee=payee,
                            txn_id=new_id,
                            amount=amount,
                            date=str(date),
                            description=description,
                            category=final_category
                        )

                    st.success(f"✅ Successfully submitted! Tracked as {new_id}")

# ==========================================
# ADMIN TABS (ONLY RUN IF LOGGED IN)
# ==========================================
if IS_TREASURER:
    # ==========================================
    # TAB 2: MONTHLY RECONCILIATION
    # ==========================================
    with tab2:
        st.header("Monthly Ledger Generator")
        records = sheet.get_all_records()
        if records:
            df = pd.DataFrame(records)
            df['Date'] = pd.to_datetime(df['Date'], errors='coerce')
            valid_dates = df.dropna(subset=['Date'])
            months_available = valid_dates['Date'].dt.strftime('%B %Y').unique()
            
            if len(months_available) > 0:
                selected_month = st.selectbox("Select Month to Process", months_available)
                month_df = valid_dates[valid_dates['Date'].dt.strftime('%B %Y') == selected_month].copy()
                
                # FIX: Remove unapproved/rejected claims from the official SA Ledger!
                month_df = month_df[~month_df['Internal Status'].isin(['Awaiting Approval', 'Rejected'])]
                
                month_df['Clean_Income'] = pd.to_numeric(month_df['Income'].astype(str).str.replace(',', ''), errors='coerce').fillna(0)
                month_df['Clean_Expense'] = pd.to_numeric(month_df['Expense'].astype(str).str.replace(',', ''), errors='coerce').fillna(0)
                
                internal_closing_str = str(month_df.iloc[-1]['Running Balance']).replace('RM', '').replace(',', '').strip() if not month_df.empty else "0.00"
                internal_closing = float(internal_closing_str)
                
                logo_ieee_url = st.text_input("IEEE Logo URL (ending in .png/.jpg)", value="https://github.com/SeanTai123/ieee-treasurer/blob/main/ieee.png?raw=true")
                logo_sa_url = st.text_input("SA Logo URL (ending in .png/.jpg)", value="https://github.com/SeanTai123/ieee-treasurer/blob/main/SA.png?raw=true")
                
                col_sig1, col_sig2 = st.columns(2)
                with col_sig1:
                    prep_name = st.text_input("Prepared by (Name)", value="Wong Yan Jie")
                    prep_email = st.text_input("Prepared by (Email)", value="efyyw6@nottingham.edu.my")
                    prep_sig_url = st.text_input("Preparer Signature URL (Optional)", value="")
                with col_sig2:
                    ver_name = st.text_input("Verified by (Name)", value="Marcus Yew Min Jie")
                    ver_email = st.text_input("Verified by (Email)", value="ecymy2@nottingham.edu.my")
                    ver_sig_url = st.text_input("Verifier Signature URL (Optional)", value="")
                    
                sig_date = st.date_input("Signature Date").strftime('%d/%m/%Y')

                if st.button(f"Generate '{selected_month}' Tab"):
                    if month_df.empty:
                        st.warning("No approved transactions found for this month.")
                    else:
                        with st.spinner("Building SA layout..."):
                            first_txn_idx = month_df.index[0]
                            if first_txn_idx > 0:
                                prev_balance_str = str(df.iloc[first_txn_idx - 1]['Running Balance']).replace('RM', '').replace(',', '').strip()
                                opening_balance = float(prev_balance_str)
                            else:
                                first_inc = month_df.iloc[0]['Clean_Income']
                                first_exp = month_df.iloc[0]['Clean_Expense']
                                first_bal = float(str(month_df.iloc[0]['Running Balance']).replace('RM', '').replace(',', '').strip())
                                opening_balance = first_bal - first_inc + first_exp

                            ieee_formula = f'=IMAGE("{logo_ieee_url}", 1)' if logo_ieee_url else "[ IEEE Logo ]"
                            sa_formula = f'=IMAGE("{logo_sa_url}", 1)' if logo_sa_url else "[ SA Logo ]"
                            prep_sig_formula = f'=IMAGE("{prep_sig_url}", 1)' if prep_sig_url else "[ Preparer Signature ]"
                            ver_sig_formula = f'=IMAGE("{ver_sig_url}", 1)' if ver_sig_url else "[ Verifier Signature ]"

                            sa_layout = [
                                [ieee_formula, "", "", "", "", sa_formula, ""],
                                ["", "", "", "", "", "", ""],
                                ["", "", "", "", "", "", ""],
                                ["IEEE UNM Student Branch", "", "", "", "", "", ""],
                                [f"Monthly Ledger for {selected_month}", "", "", "", "", "", ""],
                                ["", "", "", "", "", "", ""],
                                ["Date", "Details", "Income", "Expenses", "OR No", "CS-PV No", "Balance"],
                                ["", "", "RM", "RM", "", "", "RM"],
                                [month_df.iloc[0]['Date'].strftime('%Y-%m-%d'), "Opening Balance", "", "", "", "", f"{opening_balance:.2f}"]
                            ]

                            for _, row in month_df.iterrows():
                                if "opening balance" in str(row['Description']).lower():
                                    continue
                                inc = f"{row['Clean_Income']:.2f}" if row['Clean_Income'] > 0 else ""
                                exp = f"{row['Clean_Expense']:.2f}" if row['Clean_Expense'] > 0 else ""
                                bal = f"{float(str(row['Running Balance']).replace('RM', '').replace(',', '').strip()):.2f}"
                                sa_layout.append([row['Date'].strftime('%Y-%m-%d'), row['Description'], inc, exp, "", row['Notes'], bal])

                            total_row_idx = len(sa_layout) + 1
                            sa_layout.append(["", "", "", "", "", "TOTAL: ", f"{internal_closing:.2f}"])

                            sa_layout.extend([
                                ["", "", "", "", "", "", ""],
                                ["", "", "", "", "", "", ""],
                                [prep_sig_formula, "", "", ver_sig_formula, "", "", ""], 
                                ["", "", "", "", "", "", ""],                            
                                ["", "", "", "", "", "", ""],                            
                                [f"Prepared by: {prep_name}", "", "", f"Verified by: {ver_name}", "", "", ""],
                                [f"Email Username: {prep_email}", "", "", f"Email Username: {ver_email}", "", "", ""],
                                [f"Date: {sig_date}", "", "", f"Date: {sig_date}", "", "", ""]
                            ])

                            sheet_title = f"Ledger {selected_month}"
                            try:
                                target_sheet = client.open_by_url(SHEET_URL).worksheet(sheet_title)
                                target_sheet.clear()
                            except:
                                target_sheet = client.open_by_url(SHEET_URL).add_worksheet(title=sheet_title, rows=len(sa_layout)+5, cols=8)

                            target_sheet.update(range_name='A1', values=sa_layout, value_input_option='USER_ENTERED')
                            
                            target_sheet.merge_cells('A1:B3')
                            target_sheet.merge_cells('F1:G3')
                            target_sheet.merge_cells('A4:G4')
                            target_sheet.merge_cells('A5:G5')
                            target_sheet.format('A4:A5', {'horizontalAlignment': 'CENTER', 'textFormat': {'bold': True, 'fontSize': 12}})
                            target_sheet.merge_cells('A7:A8') 
                            target_sheet.merge_cells('B7:B8') 
                            target_sheet.merge_cells('E7:E8') 
                            target_sheet.merge_cells('F7:F8') 
                            target_sheet.format('A7:G8', {'horizontalAlignment': 'CENTER', 'verticalAlignment': 'MIDDLE', 'textFormat': {'bold': True}})
                            target_sheet.format(f'A{total_row_idx}:G{total_row_idx}', {'textFormat': {'bold': True}})
                            
                            solid_border = {"style": "SOLID", "color": {"red": 0, "green": 0, "blue": 0}}
                            target_sheet.format(f'A7:G{total_row_idx}', {"borders": {"top": solid_border, "bottom": solid_border, "left": solid_border, "right": solid_border}})
                            
                            sig_start_row = len(sa_layout) - 5 
                            target_sheet.merge_cells(f'A{sig_start_row}:B{sig_start_row+2}') 
                            target_sheet.merge_cells(f'D{sig_start_row}:E{sig_start_row+2}')

                            st.success(f"✅ Generated '{sheet_title}'!")

    # --- TAB 3: WORD DOCUMENT GENERATOR ---
    with tab3:
        st.header("📄 Payment Receipt Tracker Generator")
        st.markdown("This will automatically draw the table, pull the receipt images from ImgBB, and assemble the final Word document.")
        
        if records:
            valid_dates_tab3 = df.dropna(subset=['Date'])
            months_tab3 = valid_dates_tab3['Date'].dt.strftime('%B %Y').unique()
            
            if len(months_tab3) > 0:
                doc_month = st.selectbox("Select Month for Word Doc", months_tab3, key="doc_month")
                
                if st.button("Generate Word Document"):
                    with st.spinner("Downloading images and building document from scratch... this may take a minute."):
                        try:
                            doc = Document()
                            doc.add_heading("Payment Receipt Tracker", level=1)
                            doc.add_paragraph(f"Monthly Ledger: {doc_month}\n")
                            
                            month_data = valid_dates_tab3[valid_dates_tab3['Date'].dt.strftime('%B %Y') == doc_month].copy()
                            
                            # FIX: Remove unapproved/rejected claims from the Word Tracker!
                            month_data = month_data[~month_data['Internal Status'].isin(['Awaiting Approval', 'Rejected'])]
                            
                            month_data['Clean_Expense'] = pd.to_numeric(month_data['Expense'].astype(str).str.replace(r'[^\d.]', '', regex=True), errors='coerce').fillna(0)
                            expenses = month_data[month_data['Clean_Expense'] > 0]
                            
                            if expenses.empty:
                                st.warning("No approved expenses found for this month.")
                            else:
                                table = doc.add_table(rows=1, cols=4)
                                table.style = 'Table Grid'
                                
                                hdr_cells = table.rows[0].cells
                                hdr_cells[0].text = 'No'
                                hdr_cells[1].text = 'Receipt'
                                hdr_cells[2].text = 'Name and Amount'
                                hdr_cells[3].text = 'Claim'
                                
                                counter = 1
                                for _, row in expenses.iterrows():
                                    row_cells = table.add_row().cells
                                    row_cells[0].text = str(counter)
                                    
                                    file_url = str(row.get('Receipt/Invoice No.', row.get('Receipt Proof', '')))
                                    img_paragraph = row_cells[1].paragraphs[0]
                                    
                                    if file_url and file_url.startswith("http"):
                                        try:
                                            img_response = requests.get(file_url)
                                            if img_response.status_code == 200:
                                                fh = io.BytesIO(img_response.content)
                                                img_paragraph.add_run().add_picture(fh, width=Mm(55))
                                            else:
                                                img_paragraph.text = "[Image Error]"
                                        except Exception as e:
                                            img_paragraph.text = "[Image Download Error]"
                                    else:
                                        img_paragraph.text = "[No Image Attached]"
                                        
                                    row_cells[2].text = f"{row['Payee/Payer']}\nRM {row['Clean_Expense']:.2f}\n({row['Description']})"
                                    status = "Refunded to purchaser" if "Cleared" in str(row['Internal Status']) else "Pending PV Submission"
                                    row_cells[3].text = status
                                    counter += 1
                                
                                temp_file_name = f"temp_{doc_month.replace(' ', '_')}.docx"
                                doc.save(temp_file_name)
                                
                                with open(temp_file_name, "rb") as f:
                                    st.session_state['ready_doc'] = f.read()
                                    
                                st.session_state['ready_month'] = doc_month
                                st.success("✅ Document generated successfully! Click below to download.")
                            
                        except Exception as e:
                            st.error(f"Failed to generate document. Error: {e}")
                
                if 'ready_doc' in st.session_state and st.session_state.get('ready_month') == doc_month:
                    st.download_button(
                        label="⬇️ Download Payment Tracker (.docx)",
                        data=st.session_state['ready_doc'],
                        file_name=f"Payment_Tracker_{doc_month.replace(' ', '_')}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        type="primary"
                    )

    # --- TAB 4: FINANCIAL DASHBOARD ---
    with tab4:
        st.header("📈 Financial Dashboard")
        st.markdown("Live overview of the club's financial health.")
        
        if records:
            dash_df = pd.DataFrame(records)
            dash_df['Date'] = pd.to_datetime(dash_df['Date'], errors='coerce')
            
            # FIX: Remove unapproved and rejected claims from the dashboard charts!
            dash_df = dash_df[~dash_df['Internal Status'].isin(['Awaiting Approval', 'Rejected'])]
            
            dash_df['Clean_Income'] = pd.to_numeric(dash_df['Income'].astype(str).str.replace(r'[^\d.]', '', regex=True), errors='coerce').fillna(0)
            dash_df['Clean_Expense'] = pd.to_numeric(dash_df['Expense'].astype(str).str.replace(r'[^\d.]', '', regex=True), errors='coerce').fillna(0)
            
            real_income_df = dash_df[(dash_df['Clean_Income'] > 0) & (~dash_df['Description'].str.contains('Opening Balance', case=False, na=False))]
            
            total_real_income = real_income_df['Clean_Income'].sum()
            total_expense = dash_df['Clean_Expense'].sum()
            
            # We use the raw records for current balance to ensure math aligns with the sheet
            raw_df = pd.DataFrame(records)
            current_balance_str = str(raw_df.iloc[-1]['Running Balance']).replace('RM', '').replace(',', '').strip() if not raw_df.empty else "0.00"
            current_balance = float(current_balance_str) if current_balance_str else 0.0
            
            col_m1, col_m2, col_m3 = st.columns(3)
            col_m1.metric("💰 Current Balance", f"RM {current_balance:,.2f}")
            col_m2.metric("🟢 Total Revenue Generated", f"RM {total_real_income:,.2f}")
            col_m3.metric("🔴 Total Approved Expenses", f"RM {total_expense:,.2f}")
            
            st.divider()
            
            col_chart1, col_chart2 = st.columns(2)
            
            with col_chart1:
                st.subheader("Approved Expenses by Category")
                expense_df = dash_df[dash_df['Clean_Expense'] > 0]
                if not expense_df.empty:
                    cat_expense = expense_df.groupby('Category')['Clean_Expense'].sum().reset_index()
                    fig_pie = px.pie(cat_expense, values='Clean_Expense', names='Category', hole=0.4, 
                                     color_discrete_sequence=px.colors.sequential.RdBu)
                    st.plotly_chart(fig_pie, use_container_width=True)
                else:
                    st.info("No approved expenses recorded yet.")
                    
            with col_chart2:
                st.subheader("Monthly Cashflow")
                valid_dash_dates = dash_df.dropna(subset=['Date']).copy()
                if not valid_dash_dates.empty:
                    valid_dash_dates['Month'] = valid_dash_dates['Date'].dt.strftime('%b %Y')
                    
                    monthly_exp = valid_dash_dates.groupby('Month')['Clean_Expense'].sum().reset_index()
                    monthly_inc = real_income_df.copy()
                    monthly_inc['Month'] = monthly_inc['Date'].dt.strftime('%b %Y')
                    monthly_inc = monthly_inc.groupby('Month')['Clean_Income'].sum().reset_index()
                    
                    monthly_merged = pd.merge(monthly_exp, monthly_inc, on='Month', how='outer').fillna(0)
                    
                    if not monthly_merged.empty:
                        fig_bar = px.bar(monthly_merged, x='Month', y=['Clean_Income', 'Clean_Expense'], 
                                         barmode='group', 
                                         labels={'value': 'Amount (RM)', 'variable': 'Cashflow Type'},
                                         color_discrete_map={'Clean_Income': '#2ECC71', 'Clean_Expense': '#E74C3C'})
                        
                        newnames = {'Clean_Income': 'Income', 'Clean_Expense': 'Expense'}
                        fig_bar.for_each_trace(lambda t: t.update(name = newnames[t.name], legendgroup = newnames[t.name], hovertemplate = t.hovertemplate.replace(t.name, newnames[t.name])))
                                                                  
                        st.plotly_chart(fig_bar, use_container_width=True)
                    else:
                        st.info("Not enough data for cashflow chart.")
                else:
                    st.info("No dated transactions recorded yet.")

    # --- TAB 5: PENDING CLAIMS ---
    with tab5:
        st.header("⏳ Pending Reimbursements")
        st.markdown("Overview of all committee members waiting for claim disbursements.")

        if records:
            pending_df = pd.DataFrame(records)
            pending_df['Clean_Expense'] = pd.to_numeric(pending_df['Expense'].astype(str).str.replace(r'[^\d.]', '', regex=True), errors='coerce').fillna(0)
            
            unpaid_claims = pending_df[(pending_df['Clean_Expense'] > 0) & (~pending_df['Internal Status'].str.contains('Cleared', case=False, na=False))]
            
            if not unpaid_claims.empty:
                st.subheader("Summary by Claimant")
                summary_df = unpaid_claims.groupby('Payee/Payer')['Clean_Expense'].sum().reset_index()
                summary_df.rename(columns={'Clean_Expense': 'Total Owed (RM)', 'Payee/Payer': 'Claimant Name'}, inplace=True)
                
                summary_df['Total Owed (RM)'] = summary_df['Total Owed (RM)'].apply(lambda x: f"RM {x:,.2f}")
                st.dataframe(summary_df, use_container_width=True, hide_index=True)
                
                st.divider()
                
                st.subheader("Detailed Pending Transactions")
                display_cols = ['Date', 'Transaction ID', 'Payee/Payer', 'Description', 'Category', 'Expense', 'Internal Status']
                st.dataframe(unpaid_claims[display_cols].sort_values(by='Date', ascending=False), use_container_width=True, hide_index=True)

                st.divider()
                st.subheader("✅ Mark Claims as Cleared")

                txn_ids = unpaid_claims['Transaction ID'].tolist()
                selected_txns = st.multiselect("Select Transaction ID(s) to mark as Cleared", txn_ids)

                col_btn1, col_btn2 = st.columns([1, 4])
                with col_btn1:
                    if st.button("✅ Mark as Cleared", type="primary", disabled=len(selected_txns) == 0):
                        with st.spinner("Updating records..."):
                            all_records = sheet.get_all_records()
                            header = sheet.row_values(1)
                            status_col = header.index("Internal Status") + 1  

                            updated = 0
                            for i, record in enumerate(all_records):
                                if record.get("Transaction ID") in selected_txns:
                                    row_number = i + 2  
                                    sheet.update_cell(row_number, status_col, "Cleared")
                                    updated += 1

                            st.success(f"✅ Marked {updated} transaction(s) as Cleared!")
                            st.rerun()
            else:
                st.success("🎉 All claims have been cleared! There are no pending reimbursements.")
        else:
            st.info("No records found in the Master Ledger.")

    # --- TAB 6: MANAGE TRANSACTIONS ---
    with tab6:
        st.header("🛠️ Edit / Delete Transactions")
        st.warning("⚠️ Changes here directly modify the Master Ledger. Use with care.")

        if records:
            manage_df = pd.DataFrame(records)

            search_query = st.text_input("Search by Transaction ID or Payee Name")
            if search_query:
                filtered = manage_df[
                    manage_df['Transaction ID'].str.contains(search_query, case=False, na=False) |
                    manage_df['Payee/Payer'].str.contains(search_query, case=False, na=False)
                ]
            else:
                filtered = manage_df.tail(20)  

            st.dataframe(filtered[['Transaction ID', 'Date', 'Payee/Payer', 'Description', 'Income', 'Expense', 'Internal Status']], 
                         use_container_width=True, hide_index=True)

            st.divider()

            selected_id = st.selectbox("Select Transaction ID to Edit or Delete", filtered['Transaction ID'].tolist())

            if selected_id:
                row_data = manage_df[manage_df['Transaction ID'] == selected_id].iloc[0]
                row_index = manage_df[manage_df['Transaction ID'] == selected_id].index[0] + 2  

                with st.expander(f"✏️ Edit Transaction: {selected_id}", expanded=True):
                    new_date        = st.date_input("Date", value=pd.to_datetime(row_data['Date']))
                    new_desc        = st.text_input("Description", value=row_data['Description'])
                    new_category    = st.text_input("Category", value=row_data['Category'])
                    new_payee       = st.text_input("Payee/Payer", value=row_data['Payee/Payer'])
                    
                    status_options = ["Awaiting Approval", "Pending PV Submission", "Cleared", "Rejected"]
                    current_status = row_data.get('Internal Status', 'Awaiting Approval')
                    status_idx = status_options.index(current_status) if current_status in status_options else 0
                    
                    new_status = st.selectbox("Internal Status", status_options, index=status_idx)

                    col_save, col_delete = st.columns([1, 1])

                    with col_save:
                        if st.button("💾 Save Changes", type="primary"):
                            header = sheet.row_values(1)
                            updates = {
                                'Date': str(new_date),
                                'Description': new_desc,
                                'Category': new_category,
                                'Payee/Payer': new_payee,
                                'Internal Status': new_status,
                            }
                            for col_name, new_value in updates.items():
                                if col_name in header:
                                    col_idx = header.index(col_name) + 1
                                    sheet.update_cell(row_index, col_idx, new_value)
                            st.success(f"✅ Transaction {selected_id} updated!")
                            st.rerun()

                    with col_delete:
                        if st.button("🗑️ Delete Transaction", type="secondary"):
                            sheet.delete_rows(row_index)
                            st.warning(f"🗑️ Transaction {selected_id} deleted.")
                            st.rerun()
        else:
            st.info("No records found.")

    # --- TAB 7: APPROVAL WORKFLOW ---
    with tab7:
        st.header("🔍 Claim Approvals")
        st.markdown("Review and approve or reject newly submitted expense claims.")

        if records:
            approval_df = pd.DataFrame(records)
            approval_df['Clean_Expense'] = pd.to_numeric(
                approval_df['Expense'].astype(str).str.replace(r'[^\d.]', '', regex=True), errors='coerce'
            ).fillna(0)

            awaiting = approval_df[approval_df['Internal Status'] == "Awaiting Approval"]

            if not awaiting.empty:
                st.info(f"📋 {len(awaiting)} claim(s) awaiting your review.")

                for _, row in awaiting.iterrows():
                    with st.expander(f"🔔 {row['Transaction ID']} — {row['Payee/Payer']} — RM {row['Clean_Expense']:.2f}"):
                        st.write(f"**Date:** {row['Date']}")
                        st.write(f"**Description:** {row['Description']}")
                        st.write(f"**Category:** {row['Category']}")
                        st.write(f"**Amount:** RM {row['Clean_Expense']:.2f}")

                        receipt_url = str(row.get('Receipt/Invoice No.', ''))
                        if receipt_url.startswith("http"):
                            st.markdown(f"🧾 [View Receipt]({receipt_url})")
                        else:
                            st.write("🧾 No receipt uploaded.")

                        col_approve, col_reject = st.columns(2)

                        row_idx = approval_df[approval_df['Transaction ID'] == row['Transaction ID']].index[0] + 2
                        header  = sheet.row_values(1)
                        status_col = header.index("Internal Status") + 1

                        with col_approve:
                            if st.button(f"✅ Approve {row['Transaction ID']}", key=f"approve_{row['Transaction ID']}"):
                                sheet.update_cell(row_idx, status_col, "Pending PV Submission")
                                st.success(f"Approved {row['Transaction ID']}!")
                                st.rerun()

                        with col_reject:
                            reject_reason = st.text_input("Reason (optional)", key=f"reason_{row['Transaction ID']}")
                            if st.button(f"❌ Reject {row['Transaction ID']}", key=f"reject_{row['Transaction ID']}"):
                                notes_col = header.index("Notes") + 1 if "Notes" in header else None
                                sheet.update_cell(row_idx, status_col, "Rejected")
                                if notes_col and reject_reason:
                                    sheet.update_cell(row_idx, notes_col, f"Rejected: {reject_reason}")
                                st.warning(f"Rejected {row['Transaction ID']}.")
                                st.rerun()
            else:
                st.success("🎉 No claims awaiting approval!")
        else:
            st.info("No records found.")